import os
from datetime import datetime
from django.utils.text import slugify
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
TITLE = (
    "Таблица 1 - Перечень и характеристика автомобильных дорог, "
    "проходящих по территории муниципального округа"
)


def build_word_report(df, total_length: float, summary: str = "") -> Document:
    """
    Собирает Word Document (в памяти) по данным df.
    """
    document = Document()

    document.add_paragraph("")  # небольшой отступ строкой

    # AI-описание 
    if summary:
        p = document.add_paragraph(summary)
        document.add_paragraph("")

    # Заголовок: по центру, чуть крупнее
    title_p = document.add_paragraph()
    title_run = title_p.add_run(TITLE)
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("")  # небольшой отступ строкой

    # Таблица
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"

    # Шапка таблицы: жирная, по центру
    header_cells = table.rows[0].cells
    for col_num, column_title in enumerate(df.columns):
        p = header_cells[col_num].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(column_title))
        r.bold = True
        r.font.size = Pt(10)

    # Данные таблицы
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for col_num, cell in enumerate(row):
            text = "" if cell is None else str(cell)
            p = row_cells[col_num].paragraphs[0]
            p.clear()  # на всякий случай очищаем дефолтный пустой run
            run = p.add_run(text)
            run.font.size = Pt(10)

            # Выравнивание: номер/длина по центру, текст слева
            col_name = df.columns[col_num]
            if col_name in ("№ п/п", "Протяженность, км"):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Итоговая строка: отступ сверху + жирное значение
    document.add_paragraph("")  # отступ
    total_p = document.add_paragraph("Общая протяженность автомобильных дорог составляет ")
    total_run = total_p.add_run(f"{total_length} км")
    total_run.bold = True

    return document


def save_document(document: Document, media_root: str, file_name: str) -> str:
    """
    Сохраняет документ в media_root/file_name и возвращает полный путь.
    """
    os.makedirs(media_root, exist_ok=True)
    file_path = os.path.join(media_root, file_name)
    document.save(file_path)
    return file_path


def generate_report_filename(original_filename: str) -> str:
    """
    Генерирует имя отчета на основе имени загруженного файла.
    Пример: dorogi_report_20260303_182455.docx
    """

    # убираем расширение
    base_name = os.path.splitext(original_filename)[0]

    # делаем безопасным
    safe_name = slugify(base_name)

    # если slugify вернул пусто (редкий случай)
    if not safe_name:
        safe_name = "road_report"

    # добавляем timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    return f"{safe_name}_report_{timestamp}.docx"